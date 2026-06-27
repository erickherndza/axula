# -*- coding: utf-8 -*-
"""
core/rag.py — Motor RAG (Retrieval-Augmented Generation) para Axula

Flujo:
  1. Subida de PDF/DOCX  → extraer_texto()
  2. Texto completo      → chunkear_texto()   → lista de chunks
  3. Chunks guardados en normativa_chunks
  4. Al generar planificación → buscar_chunks() devuelve los más relevantes
  5. Chunks inyectados en el system prompt de Groq

Sin embeddings vectoriales — búsqueda por coincidencia de términos (TF simple).
Suficiente para el volumen de documentos de un centro educativo (< 50 docs).
"""

from __future__ import annotations

import io
import json
import logging
import math
import re
import unicodedata
from collections import Counter
from typing import List, Tuple

logger = logging.getLogger("axula")

# ─────────────────────────────────────────────
#  1. EXTRACCIÓN DE TEXTO
# ─────────────────────────────────────────────

def extraer_texto_pdf(ruta_o_bytes) -> str:
    """Extrae texto de un PDF. Acepta ruta de archivo o bytes."""
    try:
        import PyPDF2
        if isinstance(ruta_o_bytes, (str, bytes)) and not isinstance(ruta_o_bytes, bytes):
            f = open(ruta_o_bytes, "rb")
            close = True
        else:
            f = io.BytesIO(ruta_o_bytes)
            close = False
        reader = PyPDF2.PdfReader(f)
        partes = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                partes.append(t)
        if close:
            f.close()
        return "\n".join(partes)
    except Exception as e:
        logger.error(f"[RAG] Error extrayendo PDF: {e}")
        return ""


def extraer_texto_docx(ruta_o_bytes) -> str:
    """Extrae texto de un DOCX."""
    try:
        import docx
        if isinstance(ruta_o_bytes, bytes):
            doc = docx.Document(io.BytesIO(ruta_o_bytes))
        else:
            doc = docx.Document(ruta_o_bytes)
        partes = []
        for para in doc.paragraphs:
            if para.text.strip():
                partes.append(para.text)
        return "\n".join(partes)
    except Exception as e:
        logger.error(f"[RAG] Error extrayendo DOCX: {e}")
        return ""


def extraer_texto(ruta: str) -> str:
    """Detecta tipo por extensión y extrae texto."""
    ext = ruta.lower().split(".")[-1]
    if ext == "pdf":
        return extraer_texto_pdf(ruta)
    elif ext in ("docx", "doc"):
        return extraer_texto_docx(ruta)
    return ""


# ─────────────────────────────────────────────
#  2. CHUNKING
# ─────────────────────────────────────────────

_RE_TITULO = re.compile(
    r"^(art[ií]culo\s+\d+|cap[ií]tulo\s+\w+|secci[oó]n\s+\w+|\d+[\.\-]\s+[A-ZÁÉÍÓÚÑ]|\d+\.\d+)",
    re.IGNORECASE | re.MULTILINE
)

def chunkear_texto(texto: str, max_chars: int = 900, overlap: int = 100) -> List[dict]:
    """
    Divide el texto en chunks de ~max_chars caracteres con overlap.
    Intenta respetar límites de párrafo y detectar títulos de sección.

    Devuelve lista de dicts:
      { chunk_index, titulo_seccion, contenido, palabras_clave, longitud_chars }
    """
    if not texto.strip():
        return []

    # Normalizar saltos de línea
    texto = re.sub(r"\r\n", "\n", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)

    # Separar por párrafos dobles primero; si quedan párrafos gigantes,
    # subdividir por salto simple (común en PDFs escaneados)
    bloques_raw = [p.strip() for p in texto.split("\n\n") if p.strip()]
    parrafos = []
    for bloque in bloques_raw:
        if len(bloque) > max_chars:
            # Subdividir por salto simple
            sub = [s.strip() for s in bloque.split("\n") if s.strip()]
            parrafos.extend(sub)
        else:
            parrafos.append(bloque)

    chunks = []
    buffer = ""
    titulo_actual = None
    idx = 0

    for parrafo in parrafos:
        # Detectar si este párrafo es un título de sección
        es_titulo = bool(_RE_TITULO.match(parrafo)) and len(parrafo) < 120

        if es_titulo:
            # Guardar buffer acumulado antes de nuevo título
            if buffer.strip():
                chunks.append(_make_chunk(idx, titulo_actual, buffer))
                idx += 1
                buffer = ""
            titulo_actual = parrafo
            buffer = parrafo + "\n"
            continue

        # ¿El buffer + párrafo excede el límite?
        if len(buffer) + len(parrafo) > max_chars and buffer.strip():
            chunks.append(_make_chunk(idx, titulo_actual, buffer))
            idx += 1
            # Overlap: comenzar nuevo buffer con último fragmento
            buffer = buffer[-overlap:] + "\n" + parrafo + "\n"
        else:
            buffer += parrafo + "\n"

    # Último chunk
    if buffer.strip():
        chunks.append(_make_chunk(idx, titulo_actual, buffer))

    return chunks


def _make_chunk(idx: int, titulo: str | None, contenido: str) -> dict:
    contenido = contenido.strip()
    return {
        "chunk_index":    idx,
        "titulo_seccion": titulo or "",
        "contenido":      contenido,
        "palabras_clave": json.dumps(_extraer_keywords(contenido), ensure_ascii=False),
        "longitud_chars": len(contenido),
    }


# ─────────────────────────────────────────────
#  3. EXTRACCIÓN DE KEYWORDS (TF simple)
# ─────────────────────────────────────────────

_STOPWORDS = {
    "de","la","el","en","y","a","los","las","del","se","un","una","por","con",
    "que","al","es","su","sus","para","no","lo","le","más","pero","como","una",
    "este","esta","estos","estas","entre","sobre","según","cuando","también",
    "cada","todo","todos","toda","todas","o","e","ni","si","ha","han","ser",
    "son","fue","era","sido","tiene","tener","debe","deben","podrá","podrán",
    "dicho","dicha","dichos","dichas","artículo","artículos","inciso","párrafo",
}

def _normalizar(palabra: str) -> str:
    # Quitar acentos y pasar a minúsculas
    nfkd = unicodedata.normalize("NFKD", palabra)
    return "".join(c for c in nfkd if not unicodedata.combining(c)).lower()

def _extraer_keywords(texto: str, top_n: int = 15) -> list:
    tokens = re.findall(r"[a-záéíóúüñ]{4,}", _normalizar(texto))
    filtrados = [t for t in tokens if t not in _STOPWORDS]
    conteo = Counter(filtrados)
    return [w for w, _ in conteo.most_common(top_n)]


# ─────────────────────────────────────────────
#  4. BÚSQUEDA POR RELEVANCIA (TF simple)
# ─────────────────────────────────────────────

def buscar_chunks(
    db,
    consulta: str,
    top_k: int = 4,
    max_chars_total: int = 3000,
    mencion: str | None = None,
) -> str:
    """
    Busca los chunks más relevantes para una consulta y devuelve texto listo
    para inyectar en el system prompt de Groq.

    Scoring: coincidencias de keywords normalizados entre consulta y chunk.

    Args:
        mencion: Si se provee, filtra chunks del documento cuyo título contenga
                 la mención (ej. 'Multimedia', 'Música', 'Teatro', 'Artes Visuales').
                 Esto asegura que el contexto RAG sea específico de la mención activa.
    """
    keywords_consulta = set(_extraer_keywords(consulta, top_n=20))
    if not keywords_consulta:
        return ""

    try:
        if mencion:
            cur = db.execute("""
                SELECT nc.contenido, nc.palabras_clave, nc.titulo_seccion,
                       nm.titulo, nm.tipo, nm.anio_escolar
                FROM normativa_chunks nc
                JOIN normativa_minerd nm ON nm.id = nc.normativa_id
                WHERE nm.activo = 1 AND nm.procesado = 1
                  AND nm.titulo LIKE ?
            """, (f"%{mencion}%",))
        else:
            cur = db.execute("""
                SELECT nc.contenido, nc.palabras_clave, nc.titulo_seccion,
                       nm.titulo, nm.tipo, nm.anio_escolar
                FROM normativa_chunks nc
                JOIN normativa_minerd nm ON nm.id = nc.normativa_id
                WHERE nm.activo = 1 AND nm.procesado = 1
            """)
        filas = cur.fetchall()
    except Exception as e:
        logger.error(f"[RAG] Error consultando chunks: {e}")
        return ""

    if not filas:
        return ""

    scored: List[Tuple[float, dict]] = []
    for contenido, palabras_json, titulo_sec, titulo_doc, tipo_doc, anio in filas:
        try:
            kw_chunk = set(json.loads(palabras_json or "[]"))
        except Exception:
            kw_chunk = set()
        interseccion = len(keywords_consulta & kw_chunk)
        if interseccion == 0:
            continue
        # Boost si hay coincidencia directa de texto
        boost = 0
        consulta_norm = _normalizar(consulta)
        for kw in keywords_consulta:
            if kw in _normalizar(contenido):
                boost += 1
        score = interseccion + boost * 0.5
        scored.append((score, {
            "contenido":   contenido,
            "titulo_sec":  titulo_sec or "",
            "titulo_doc":  titulo_doc,
            "tipo_doc":    tipo_doc,
            "anio":        anio or "",
        }))

    if not scored:
        return ""

    scored.sort(key=lambda x: x[0], reverse=True)
    top = scored[:top_k]

    # Construir bloque de contexto para el prompt
    MAX_CHARS_POR_CHUNK = 800   # Truncar chunks individuales muy largos
    partes = ["=== NORMATIVA MINERD VIGENTE (contexto oficial) ==="]
    chars_total = 0
    for _, chunk in top:
        etiqueta = f"[{chunk['tipo_doc'].upper()} {chunk['anio']}]"
        if chunk['titulo_sec']:
            etiqueta += f" {chunk['titulo_sec']}"
        contenido_truncado = chunk['contenido'][:MAX_CHARS_POR_CHUNK]
        fragmento = f"{etiqueta}\n{contenido_truncado}"
        if chars_total + len(fragmento) > max_chars_total:
            break
        partes.append(fragmento)
        chars_total += len(fragmento)

    if len(partes) <= 1:
        return ""

    partes.append("=== FIN NORMATIVA ===")
    return "\n\n".join(partes)


# ─────────────────────────────────────────────
#  5. PROCESADO COMPLETO DE UN DOCUMENTO
# ─────────────────────────────────────────────

def procesar_normativa(db, normativa_id: int, ruta_archivo: str) -> bool:
    """
    Extrae texto, chunkea, guarda en normativa_chunks y actualiza el registro.
    Llamar después de subir el archivo.
    """
    try:
        texto = extraer_texto(ruta_archivo)
        if not texto.strip():
            db.execute(
                "UPDATE normativa_minerd SET procesado=2 WHERE id=?",
                (normativa_id,)
            )
            db.commit()
            logger.warning(f"[RAG] normativa {normativa_id}: sin texto extraíble")
            return False

        chunks = chunkear_texto(texto)
        if not chunks:
            db.execute(
                "UPDATE normativa_minerd SET procesado=2 WHERE id=?",
                (normativa_id,)
            )
            db.commit()
            return False

        # Eliminar chunks previos (re-procesado)
        db.execute("DELETE FROM normativa_chunks WHERE normativa_id=?", (normativa_id,))

        for ch in chunks:
            db.execute("""
                INSERT INTO normativa_chunks
                    (normativa_id, chunk_index, titulo_seccion,
                     contenido, palabras_clave, longitud_chars)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                normativa_id,
                ch["chunk_index"],
                ch["titulo_seccion"],
                ch["contenido"],
                ch["palabras_clave"],
                ch["longitud_chars"],
            ))

        db.execute("""
            UPDATE normativa_minerd
            SET procesado=1, total_chunks=?
            WHERE id=?
        """, (len(chunks), normativa_id))
        db.commit()

        logger.info(f"[RAG] normativa {normativa_id}: {len(chunks)} chunks procesados")
        return True

    except Exception as e:
        logger.error(f"[RAG] Error procesando normativa {normativa_id}: {e}")
        try:
            db.execute("UPDATE normativa_minerd SET procesado=2 WHERE id=?", (normativa_id,))
            db.commit()
        except Exception:
            pass
        return False
