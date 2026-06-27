# -*- coding: utf-8 -*-
"""Blueprint: planificacion_basica — Planificación Secuencial para Materias Básicas."""

import sqlite3
import json
import logging
from datetime import datetime
from io import BytesIO

from flask import (
    Blueprint, render_template, request, jsonify,
    redirect, url_for, session, send_file,
)

from core.constants import DATABASE, ROLES_COORD, ROLES_DIRECTORA
from core.database import get_db
from core.auth import login_required, get_usuario, _csrf_check
from core.helpers import _anio_escolar_actual, _normalizar_rol

logger = logging.getLogger("axula")

planificacion_basica_bp = Blueprint("planificacion_basica_bp", __name__)


def _sanitizar_control_chars(s: str) -> str:
    """
    Reemplaza caracteres de control literales (\\n, \\r, \\t) dentro de
    strings JSON con sus secuencias de escape, evitando que json.loads falle.
    Recorre carácter a carácter rastreando si estamos dentro de un string.
    """
    result = []
    in_string = False
    escape_next = False
    for ch in s:
        if escape_next:
            result.append(ch)
            escape_next = False
        elif ch == "\\":
            result.append(ch)
            escape_next = True
        elif ch == '"':
            result.append(ch)
            in_string = not in_string
        elif in_string and ch == "\n":
            result.append("\\n")
        elif in_string and ch == "\r":
            result.append("\\r")
        elif in_string and ch == "\t":
            result.append("\\t")
        else:
            result.append(ch)
    return "".join(result)


def _extraer_json_ia(raw: str) -> dict:
    """
    Extrae y parsea el JSON de la respuesta de la IA aunque venga con:
    - Markdown ```json ... ```
    - Texto antes o después del JSON
    - Saltos de línea literales dentro de strings JSON
    """
    import re as _re

    text = raw.strip()

    # 1. Quitar bloques markdown ```json ... ``` o ``` ... ```
    text = _re.sub(r"^```(?:json)?\s*", "", text, flags=_re.IGNORECASE)
    text = _re.sub(r"\s*```\s*$", "", text)
    text = text.strip()

    # 2. Sanitizar caracteres de control dentro de strings
    text = _sanitizar_control_chars(text)

    # 3. Intentar parsear directo
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 4. Buscar el primer { ... } balanceado
    start = text.find("{")
    if start == -1:
        raise json.JSONDecodeError("No se encontró JSON en la respuesta", text, 0)

    depth = 0
    end = start
    for i, ch in enumerate(text[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i
                break

    candidate = text[start:end + 1]
    return json.loads(candidate)

# ── Roles permitidos ──────────────────────────────────────────────────────────
_ROLES_PERMITIDOS = {"profesor"} | ROLES_COORD | ROLES_DIRECTORA


def _check_acceso():
    """
    Devuelve el usuario si tiene acceso, None en caso contrario.
    Profesores: solo tipo_docencia='basica' (materias básicas 1ro-6to sec.)
    Coordinadores y directora: siempre tienen acceso (supervisión).
    """
    u = get_usuario()
    if not u:
        return None
    rol = _normalizar_rol(u.get("rol", ""))
    # Coordinadores y directora: acceso total
    if rol in ROLES_COORD | ROLES_DIRECTORA:
        return u
    # Profesores: verificar tipo_docencia
    if rol == "profesor":
        tipo = u.get("tipo_docencia", "")
        # Si no está en sesión (sesión anterior al cambio), consultar la DB
        if not tipo:
            try:
                with sqlite3.connect(DATABASE, timeout=5) as conn:
                    row = conn.execute(
                        "SELECT tipo_docencia FROM usuarios WHERE id=?", (u["id"],)
                    ).fetchone()
                    tipo = (row[0] or "") if row else ""
                    logger.info(f"[PLAN_BASICA] user_id={u['id']} tipo_docencia DB={tipo!r}")
            except Exception as e:
                logger.error(f"[PLAN_BASICA] Error consultando tipo_docencia: {e}")
                tipo = ""
        logger.info(f"[PLAN_BASICA] acceso profesor id={u['id']} tipo={tipo!r}")
        # Solo bloquear si tipo es explícitamente 'tecnica'
        # Si está vacío o desconocido, permite acceso (no penalizar sesiones antiguas)
        if tipo == "tecnica":
            return None
        return u
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# LISTADO
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica")
@login_required
def lista_planificaciones():
    u = _check_acceso()
    if not u:
        return redirect("/home")

    rol = _normalizar_rol(u.get("rol", ""))
    anio = _anio_escolar_actual()

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row

        if rol in ROLES_DIRECTORA | ROLES_COORD:
            rows = conn.execute("""
                SELECT ps.*, acb.nombre AS area_nombre, acb.codigo AS area_codigo,
                       u.nombre AS profesor_nombre
                FROM planificacion_secuencial ps
                JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
                JOIN usuarios u ON u.id = ps.profesor_id
                WHERE ps.anio_escolar = ?
                ORDER BY ps.fecha_modificacion DESC
            """, (anio,)).fetchall()
        else:
            rows = conn.execute("""
                SELECT ps.*, acb.nombre AS area_nombre, acb.codigo AS area_codigo,
                       u.nombre AS profesor_nombre
                FROM planificacion_secuencial ps
                JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
                JOIN usuarios u ON u.id = ps.profesor_id
                WHERE ps.profesor_id = ? AND ps.anio_escolar = ?
                ORDER BY ps.fecha_modificacion DESC
            """, (u["id"], anio)).fetchall()

        areas = conn.execute(
            "SELECT * FROM areas_curriculares_basicas WHERE activa=1 ORDER BY nombre"
        ).fetchall()

        plantillas = conn.execute("""
            SELECT ps.*, acb.nombre AS area_nombre, acb.codigo AS area_codigo
            FROM planificacion_secuencial ps
            JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
            WHERE ps.es_plantilla = 1
            ORDER BY acb.nombre, ps.grado, ps.periodo
        """).fetchall()

    return render_template(
        "planificacion_basica.html",
        current_user=u,
        planificaciones=[dict(r) for r in rows],
        plantillas=[dict(r) for r in plantillas],
        areas=[dict(a) for a in areas],
        anio=anio,
        vista="lista",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# CREAR NUEVA PLANIFICACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/nueva", methods=["POST"])
@login_required
def nueva_planificacion():
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()

    area_id  = request.form.get("area_id", type=int)
    grado    = request.form.get("grado", "").strip()
    periodo  = request.form.get("periodo", type=int)
    salida   = request.form.get("salida_optativa", "").strip() or None
    anio     = _anio_escolar_actual()

    if not area_id or not grado or not periodo:
        return jsonify({"error": "Faltan datos obligatorios"}), 400

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        # Verificar que no exista ya esta combinación para el profesor
        existe = conn.execute("""
            SELECT id FROM planificacion_secuencial
            WHERE profesor_id=? AND area_id=? AND grado=? AND periodo=? AND anio_escolar=?
        """, (u["id"], area_id, grado, periodo, anio)).fetchone()

        if existe:
            return jsonify({"error": "Ya existe una planificación para esa área, grado y período", "id": existe[0]}), 409

        cur = conn.execute("""
            INSERT INTO planificacion_secuencial
                (profesor_id, area_id, grado, anio_escolar, periodo, salida_optativa, estado, generado_ia)
            VALUES (?,?,?,?,?,?,'borrador',0)
        """, (u["id"], area_id, grado, anio, periodo, salida))
        plan_id = cur.lastrowid
        conn.commit()

    return jsonify({"ok": True, "id": plan_id})


# ═══════════════════════════════════════════════════════════════════════════════
# DETALLE / EDITOR
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>")
@login_required
def detalle_planificacion(plan_id):
    u = _check_acceso()
    if not u:
        return redirect("/home")

    rol = _normalizar_rol(u.get("rol", ""))

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row

        plan = conn.execute("""
            SELECT ps.*, acb.nombre AS area_nombre, acb.codigo AS area_codigo
            FROM planificacion_secuencial ps
            JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
            WHERE ps.id = ?
        """, (plan_id,)).fetchone()

        if not plan:
            return redirect("/planificacion-basica")

        # Solo el dueño o coordinadores/directora pueden ver
        if plan["profesor_id"] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return redirect("/planificacion-basica")

        unidades = conn.execute("""
            SELECT * FROM unidades_secuenciales
            WHERE planificacion_id = ?
            ORDER BY numero_unidad ASC
        """, (plan_id,)).fetchall()

        areas = conn.execute(
            "SELECT * FROM areas_curriculares_basicas WHERE activa=1 ORDER BY nombre"
        ).fetchall()

    return render_template(
        "planificacion_basica.html",
        current_user=u,
        plan=dict(plan),
        unidades=[dict(r) for r in unidades],
        areas=[dict(a) for a in areas],
        vista="editor",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# GUARDAR / EDITAR PLAN (cabecera)
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/estado", methods=["POST"])
@login_required
def cambiar_estado(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()
    estado = request.json.get("estado", "borrador")
    if estado not in ("borrador", "completo", "aprobado"):
        return jsonify({"error": "Estado inválido"}), 400

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        plan = conn.execute(
            "SELECT profesor_id FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "No encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan[0] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        conn.execute(
            "UPDATE planificacion_secuencial SET estado=?, fecha_modificacion=CURRENT_TIMESTAMP WHERE id=?",
            (estado, plan_id)
        )
        conn.commit()

    return jsonify({"ok": True, "estado": estado})


# ═══════════════════════════════════════════════════════════════════════════════
# CRUD UNIDADES
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/unidad", methods=["POST"])
@login_required
def guardar_unidad(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()
    data = request.json or {}

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        plan = conn.execute(
            "SELECT profesor_id FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "Plan no encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan[0] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        uid = data.get("id")  # Si tiene id → editar, si no → crear

        campos = (
            data.get("numero_unidad", 1),
            data.get("titulo", ""),
            data.get("semanas", 4),
            data.get("objetivos", ""),
            data.get("indicadores", ""),
            data.get("contenidos_conceptuales", ""),
            data.get("contenidos_procedimentales", ""),
            data.get("contenidos_actitudinales", ""),
            data.get("estrategias", ""),
            data.get("recursos", ""),
            data.get("tipo_evaluacion", "formativa"),
            data.get("instrumentos", ""),
            data.get("criterios", ""),
            data.get("porcentaje_calificacion"),
            json.dumps(data.get("transversalidades", {}), ensure_ascii=False),
            data.get("observaciones", ""),
            data.get("situacion_aprendizaje", ""),
            data.get("preguntas_generadoras", ""),
            data.get("competencias_especificas", ""),
            data.get("bibliografia", ""),
        )

        if uid:
            conn.execute("""
                UPDATE unidades_secuenciales SET
                    numero_unidad=?, titulo=?, semanas=?, objetivos=?, indicadores=?,
                    contenidos_conceptuales=?, contenidos_procedimentales=?, contenidos_actitudinales=?,
                    estrategias=?, recursos=?, tipo_evaluacion=?, instrumentos=?, criterios=?,
                    porcentaje_calificacion=?, transversalidades=?, observaciones=?,
                    situacion_aprendizaje=?, preguntas_generadoras=?, competencias_especificas=?, bibliografia=?
                WHERE id=? AND planificacion_id=?
            """, (*campos, uid, plan_id))
            unidad_id = uid
        else:
            cur = conn.execute("""
                INSERT INTO unidades_secuenciales
                    (planificacion_id, numero_unidad, titulo, semanas, objetivos, indicadores,
                     contenidos_conceptuales, contenidos_procedimentales, contenidos_actitudinales,
                     estrategias, recursos, tipo_evaluacion, instrumentos, criterios,
                     porcentaje_calificacion, transversalidades, observaciones,
                     situacion_aprendizaje, preguntas_generadoras, competencias_especificas, bibliografia)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (plan_id, *campos))
            unidad_id = cur.lastrowid

        conn.execute(
            "UPDATE planificacion_secuencial SET fecha_modificacion=CURRENT_TIMESTAMP WHERE id=?",
            (plan_id,)
        )
        conn.commit()

    return jsonify({"ok": True, "id": unidad_id})


@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/unidad/<int:uid>", methods=["DELETE"])
@login_required
def eliminar_unidad(plan_id, uid):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        plan = conn.execute(
            "SELECT profesor_id FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "Plan no encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan[0] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        conn.execute(
            "DELETE FROM unidades_secuenciales WHERE id=? AND planificacion_id=?",
            (uid, plan_id)
        )
        conn.commit()

    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# GENERAR CON IA
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/generar-ia", methods=["POST"])
@login_required
def generar_ia(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()
    data = request.json or {}

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row
        plan = conn.execute("""
            SELECT ps.*, acb.nombre AS area_nombre
            FROM planificacion_secuencial ps
            JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
            WHERE ps.id = ?
        """, (plan_id,)).fetchone()

        if not plan:
            return jsonify({"error": "Plan no encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan["profesor_id"] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

    from core.ia import _get_groq_client, construir_prompt_secuencial, _sanitizar_campo

    titulo       = _sanitizar_campo(data.get("titulo", "Unidad sin título"), 150)
    numero       = data.get("numero_unidad", 1)
    semanas      = data.get("semanas", 4)

    prompt = construir_prompt_secuencial(
        area_nombre   = plan["area_nombre"],
        grado         = plan["grado"],
        periodo       = plan["periodo"],
        numero_unidad = numero,
        titulo        = titulo,
        semanas       = semanas,
    )

    try:
        client = _get_groq_client()
        resp = client.chat.completions.create(
            model       = "llama-3.3-70b-versatile",
            messages    = [{"role": "user", "content": prompt}],
            temperature = 0.5,
            max_tokens  = 6000,
        )
        raw = resp.choices[0].message.content.strip()
        logger.info("IA secuencial raw (200c): %s", raw[:200])

        resultado = _extraer_json_ia(raw)

        # Marcar plan como generado con IA
        with sqlite3.connect(DATABASE, timeout=10) as conn:
            conn.execute(
                "UPDATE planificacion_secuencial SET generado_ia=1, fecha_modificacion=CURRENT_TIMESTAMP WHERE id=?",
                (plan_id,)
            )
            conn.commit()

        return jsonify({"ok": True, "data": resultado})

    except json.JSONDecodeError as jex:
        logger.error("IA JSON inválido: %s | raw: %s", jex, raw[:400])
        return jsonify({"error": "La IA no devolvió JSON válido. Intenta de nuevo.", "raw": raw[:300]}), 500
    except Exception as exc:
        logger.error("Error IA secuencial: %s", exc)
        return jsonify({"error": "Error al generar la secuencia didáctica. Intenta de nuevo."}), 500


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORTAR DOCX
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/exportar-docx")
@login_required
def exportar_docx(plan_id):
    u = _check_acceso()
    if not u:
        return redirect("/home")

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row

        plan = conn.execute("""
            SELECT ps.*, acb.nombre AS area_nombre, acb.codigo AS area_codigo,
                   usr.nombre AS profesor_nombre
            FROM planificacion_secuencial ps
            JOIN areas_curriculares_basicas acb ON acb.id = ps.area_id
            JOIN usuarios usr ON usr.id = ps.profesor_id
            WHERE ps.id = ?
        """, (plan_id,)).fetchone()

        if not plan:
            return redirect("/planificacion-basica")

        rol = _normalizar_rol(u.get("rol", ""))
        if plan["profesor_id"] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return redirect("/planificacion-basica")

        unidades = conn.execute("""
            SELECT * FROM unidades_secuenciales
            WHERE planificacion_id = ?
            ORDER BY numero_unidad ASC
        """, (plan_id,)).fetchall()

    # ── Generar DOCX ──────────────────────────────────────────────────────────
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from core.helpers import _get_config_centro

    cfg  = _get_config_centro()
    plan = dict(plan)

    TEAL      = RGBColor(0x03, 0x8C, 0x8C)
    TEAL_BG   = "038C8C"
    GRAY_BG   = "F2F2F2"
    WHITE_BG  = "FFFFFF"

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Helper: color de fondo en celda ───────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)

    # ── Helper: añadir párrafo con texto en celda (sin limpiar existentes) ────
    def cell_text(cell, text, bold=False, size=10, color=None, align=None):
        p = cell.paragraphs[0]
        p.clear()
        if align:
            p.alignment = align
        run = p.add_run(text or "")
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    # ── Helper: sección con encabezado teal ───────────────────────────────────
    def seccion_header(doc, texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(texto.upper())
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = TEAL
        # Línea bajo el encabezado
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"),  "6")
        bottom.set(qn("w:color"), TEAL_BG)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Helper: texto multilinea en un párrafo ────────────────────────────────
    def parrafo_multilinea(doc, texto, size=10, indent=False):
        if not texto or texto == "—":
            p = doc.add_paragraph("—")
            p.runs[0].font.size = Pt(size)
            return
        for linea in str(texto).split("\\n"):
            linea = linea.strip()
            if not linea:
                continue
            p = doc.add_paragraph()
            if indent and linea.startswith("•"):
                p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(linea)
            run.font.size = Pt(size)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)

    # ── Helper: lista JSON o texto como párrafos con viñeta ───────────────────
    def lista_o_texto(doc, raw, size=10):
        if not raw:
            doc.add_paragraph("—").runs[0].font.size = Pt(size)
            return
        try:
            items = json.loads(raw) if str(raw).startswith("[") else None
        except Exception:
            items = None
        if items:
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.3)
                run = p.add_run(str(item))
                run.font.size = Pt(size)
                p.paragraph_format.space_after = Pt(2)
        else:
            parrafo_multilinea(doc, raw, size=size, indent=True)

    # ══════════════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()

    # Logo (si existe)
    import os as _os
    logo_path = _os.path.join(_os.path.dirname(__file__), "..", "static", "logo-benito.jpg")
    if _os.path.isfile(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(logo_path, width=Cm(3.5))

    # Nombre del centro
    p_centro = doc.add_paragraph()
    p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_centro.add_run(cfg.get("nombre", "Centro Educativo"))
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TEAL

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Nivel Secundario — Modalidad Académica")
    r2.font.size = Pt(11)

    doc.add_paragraph()

    # Título del documento
    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_tipo.add_run("SECUENCIA DIDÁCTICA")
    r3.bold = True; r3.font.size = Pt(16); r3.font.color.rgb = TEAL

    doc.add_paragraph()

    # Tabla de identificación
    id_tbl = doc.add_table(rows=5, cols=4)
    id_tbl.style = "Table Grid"
    id_data = [
        ("Área / Materia",  plan["area_nombre"],        "Año Escolar",    plan["anio_escolar"]),
        ("Grado",           f"{plan['grado']} de Sec.", "Período",        f"P{plan['periodo']}"),
        ("Docente",         plan["profesor_nombre"],    "Secciones",      ""),
        ("Centro",          cfg.get("nombre", ""),      "Distrito",       cfg.get("distrito", "10-04")),
        ("Año Escolar",     plan["anio_escolar"],       "Regional",       cfg.get("regional", "10")),
    ]
    for row_i, (l1, v1, l2, v2) in enumerate(id_data):
        cells = id_tbl.rows[row_i].cells
        cell_text(cells[0], l1, bold=True, size=9)
        cell_text(cells[1], v1, size=9)
        cell_text(cells[2], l2, bold=True, size=9)
        cell_text(cells[3], v2, size=9)
        set_cell_bg(cells[0], GRAY_BG)
        set_cell_bg(cells[2], GRAY_BG)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # POR CADA UNIDAD
    # ══════════════════════════════════════════════════════════════════════════
    for und in unidades:
        und = dict(und)

        # Encabezado de unidad
        h = doc.add_heading(f"Unidad {und['numero_unidad']}: {und['titulo']}", level=1)
        h.runs[0].font.color.rgb = TEAL
        h.runs[0].font.size      = Pt(13)

        # Duración
        p_dur = doc.add_paragraph()
        p_dur.add_run(f"Duración: ").bold = True
        p_dur.add_run(f"{und.get('semanas', '—')} semanas")
        p_dur.paragraph_format.space_after = Pt(8)

        # ── Situación de Aprendizaje ──────────────────────────────────────────
        if und.get("situacion_aprendizaje"):
            seccion_header(doc, "Situación de Aprendizaje")
            parrafo_multilinea(doc, und["situacion_aprendizaje"])

        # ── Preguntas Generadoras ─────────────────────────────────────────────
        if und.get("preguntas_generadoras"):
            seccion_header(doc, "Preguntas Generadoras")
            lista_o_texto(doc, und["preguntas_generadoras"])

        # ── Competencias Específicas ──────────────────────────────────────────
        if und.get("competencias_especificas"):
            seccion_header(doc, "Competencias Específicas")
            parrafo_multilinea(doc, und["competencias_especificas"], indent=True)

        # ── Objetivos ─────────────────────────────────────────────────────────
        seccion_header(doc, "Objetivos de la Unidad")
        lista_o_texto(doc, und.get("objetivos", ""))

        # ── Indicadores de Logro ──────────────────────────────────────────────
        seccion_header(doc, "Indicadores de Logro")
        lista_o_texto(doc, und.get("indicadores", ""))

        # ── Contenidos ────────────────────────────────────────────────────────
        seccion_header(doc, "Contenidos")
        cont_tbl = doc.add_table(rows=2, cols=3)
        cont_tbl.style = "Table Grid"
        hdrs = ["Conceptuales", "Procedimentales", "Actitudinales"]
        vals = [
            und.get("contenidos_conceptuales", ""),
            und.get("contenidos_procedimentales", ""),
            und.get("contenidos_actitudinales", ""),
        ]
        for ci, hdr in enumerate(hdrs):
            c = cont_tbl.rows[0].cells[ci]
            set_cell_bg(c, TEAL_BG)
            cell_text(c, hdr, bold=True, size=9,
                      color=RGBColor(0xFF, 0xFF, 0xFF),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        for ci, v in enumerate(vals):
            parr = cont_tbl.rows[1].cells[ci].paragraphs[0]
            parr.clear()
            for linea in str(v or "—").split("\\n"):
                linea = linea.strip()
                if linea:
                    run = parr.add_run(linea + "\n")
                    run.font.size = Pt(9)

        doc.add_paragraph()

        # ── Estrategias por Semana ────────────────────────────────────────────
        seccion_header(doc, "Secuencia de Actividades por Semana")
        parrafo_multilinea(doc, und.get("estrategias", ""), size=10, indent=False)

        # ── Recursos ──────────────────────────────────────────────────────────
        seccion_header(doc, "Recursos y Materiales")
        parrafo_multilinea(doc, und.get("recursos", ""), indent=True)

        # ── Evaluación ────────────────────────────────────────────────────────
        seccion_header(doc, "Evaluación")

        eval_tbl = doc.add_table(rows=2, cols=3)
        eval_tbl.style = "Table Grid"
        eval_hdrs = ["Tipo de Evaluación", "Instrumentos", "Criterios de Evaluación"]
        for ci, h_txt in enumerate(eval_hdrs):
            c = eval_tbl.rows[0].cells[ci]
            set_cell_bg(c, GRAY_BG)
            cell_text(c, h_txt, bold=True, size=9)

        tipo_map = {
            "formativa":         "Formativa",
            "sumativa":          "Sumativa",
            "formativa_sumativa":"Formativa y Sumativa",
            "diagnostica":       "Diagnóstica",
        }
        tipo_txt = tipo_map.get(und.get("tipo_evaluacion", ""), und.get("tipo_evaluacion", "—"))

        eval_vals = [tipo_txt, und.get("instrumentos", ""), und.get("criterios", "")]
        for ci, v in enumerate(eval_vals):
            cell = eval_tbl.rows[1].cells[ci]
            parr = cell.paragraphs[0]
            parr.clear()
            for linea in str(v or "—").split("\\n"):
                linea = linea.strip()
                if linea:
                    run = parr.add_run(linea + "\n")
                    run.font.size = Pt(9)

        doc.add_paragraph()

        # ── Transversalidades ─────────────────────────────────────────────────
        try:
            trans = json.loads(und.get("transversalidades") or "{}")
        except Exception:
            trans = {}
        trans_labels = {
            "ambiental": "Educación Ambiental",
            "genero":    "Perspectiva de Género",
            "tics":      "Tecnologías de la Información (TIC)",
            "riesgo":    "Gestión de Riesgo",
        }
        activas = [v for k, v in trans_labels.items() if trans.get(k)]
        if activas:
            seccion_header(doc, "Ejes Transversales")
            for a in activas:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(a).font.size = Pt(10)

        # ── Porcentaje ────────────────────────────────────────────────────────
        if und.get("porcentaje_calificacion"):
            p_pct = doc.add_paragraph()
            p_pct.add_run("Porcentaje de Calificación: ").bold = True
            p_pct.add_run(f"{und['porcentaje_calificacion']}%")

        # ── Bibliografía ──────────────────────────────────────────────────────
        if und.get("bibliografia"):
            seccion_header(doc, "Bibliografía y Recursos de Apoyo")
            parrafo_multilinea(doc, und["bibliografia"], indent=True)

        # ── Observaciones ─────────────────────────────────────────────────────
        if und.get("observaciones"):
            seccion_header(doc, "Observaciones")
            doc.add_paragraph(und["observaciones"]).runs[0].font.size = Pt(10)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # FIRMAS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    seccion_header(doc, "Aprobación y Firmas")
    doc.add_paragraph()

    firma_tbl = doc.add_table(rows=4, cols=2)
    firma_tbl.style = "Table Grid"
    firma_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fila header
    set_cell_bg(firma_tbl.rows[0].cells[0], GRAY_BG)
    set_cell_bg(firma_tbl.rows[0].cells[1], GRAY_BG)
    cell_text(firma_tbl.rows[0].cells[0], "Docente", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(firma_tbl.rows[0].cells[1], "Director/a / Coordinador/a", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell_text(firma_tbl.rows[1].cells[0], " " * 50, size=10)
    cell_text(firma_tbl.rows[1].cells[1], " " * 50, size=10)
    cell_text(firma_tbl.rows[2].cells[0], plan["profesor_nombre"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(firma_tbl.rows[2].cells[1], cfg.get("directora", ""), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(firma_tbl.rows[3].cells[0], "Firma y fecha", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(firma_tbl.rows[3].cells[1], "Firma y fecha", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Enviar archivo ────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    nombre_archivo = f"SecuenciaDidactica_{plan['area_codigo']}_{plan['grado'].replace(' ','_')}_P{plan['periodo']}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=nombre_archivo,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# DUPLICAR PLANIFICACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/eliminar", methods=["DELETE"])
@login_required
def eliminar_plan(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        plan = conn.execute(
            "SELECT profesor_id FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "No encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan[0] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        conn.execute("DELETE FROM unidades_secuenciales WHERE planificacion_id=?", (plan_id,))
        conn.execute("DELETE FROM planificacion_secuencial WHERE id=?", (plan_id,))
        conn.commit()

    return jsonify({"ok": True})


@planificacion_basica_bp.route("/planificacion-basica/<int:plan_id>/duplicar", methods=["POST"])
@login_required
def duplicar_planificacion(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    _csrf_check()
    data = request.json or {}
    nuevo_anio    = data.get("anio_escolar", _anio_escolar_actual())
    nuevo_periodo = data.get("periodo")

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row
        plan = conn.execute(
            "SELECT * FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "No encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan["profesor_id"] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        cur = conn.execute("""
            INSERT INTO planificacion_secuencial
                (profesor_id, area_id, grado, anio_escolar, periodo, salida_optativa, estado, generado_ia)
            VALUES (?,?,?,?,?,?,'borrador',0)
        """, (
            plan["profesor_id"],
            plan["area_id"],
            plan["grado"],
            nuevo_anio,
            nuevo_periodo or plan["periodo"],
            plan["salida_optativa"],
        ))
        nuevo_id = cur.lastrowid

        unidades = conn.execute(
            "SELECT * FROM unidades_secuenciales WHERE planificacion_id=?", (plan_id,)
        ).fetchall()

        for und in unidades:
            conn.execute("""
                INSERT INTO unidades_secuenciales
                    (planificacion_id, numero_unidad, titulo, semanas, objetivos,
                     contenidos_conceptuales, contenidos_procedimentales, contenidos_actitudinales,
                     estrategias, recursos, tipo_evaluacion, instrumentos, criterios,
                     porcentaje_calificacion, transversalidades, observaciones)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                nuevo_id,
                und["numero_unidad"], und["titulo"], und["semanas"], und["objetivos"],
                und["contenidos_conceptuales"], und["contenidos_procedimentales"],
                und["contenidos_actitudinales"], und["estrategias"], und["recursos"],
                und["tipo_evaluacion"], und["instrumentos"], und["criterios"],
                und["porcentaje_calificacion"], und["transversalidades"], und["observaciones"],
            ))

        conn.commit()

    return jsonify({"ok": True, "id": nuevo_id})


# ═══════════════════════════════════════════════════════════════════════════════
# API — DATOS PARA JS
# ═══════════════════════════════════════════════════════════════════════════════

@planificacion_basica_bp.route("/api/planificacion-basica/<int:plan_id>/unidades")
@login_required
def api_unidades(plan_id):
    u = _check_acceso()
    if not u:
        return jsonify({"error": "Sin permisos"}), 403

    with sqlite3.connect(DATABASE, timeout=10) as conn:
        conn.row_factory = sqlite3.Row
        plan = conn.execute(
            "SELECT profesor_id FROM planificacion_secuencial WHERE id=?", (plan_id,)
        ).fetchone()
        if not plan:
            return jsonify({"error": "No encontrado"}), 404

        rol = _normalizar_rol(u.get("rol", ""))
        if plan["profesor_id"] != u["id"] and rol not in ROLES_DIRECTORA | ROLES_COORD:
            return jsonify({"error": "Sin permisos"}), 403

        unidades = conn.execute(
            "SELECT * FROM unidades_secuenciales WHERE planificacion_id=? ORDER BY numero_unidad",
            (plan_id,)
        ).fetchall()

    return jsonify([dict(r) for r in unidades])
